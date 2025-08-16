import streamlit as st
import pandas as pd
import pdfkit
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Mm
from num2words import num2words
import os
import zipfile
import tempfile
from jinja2 import Environment, FileSystemLoader, Template
from pypdf import PdfReader, PdfWriter
import numpy as np
import platform
from datetime import datetime
import subprocess
import shutil
import traceback
import argparse
import requests
from xhtml2pdf import pisa
import tarfile
import io

# Page setup and branding
_local_logo = os.path.join(os.getcwd(), "crane_rajkumar.png")
_page_icon = _local_logo if os.path.exists(_local_logo) else "📄"
st.set_page_config(page_title="Bill Generator", page_icon=_page_icon, layout="wide")

def resolve_logo_url() -> str | None:
    candidates = [
        # Prefer local crane_rajkumar.png if present
        _local_logo if os.path.exists(_local_logo) else None,
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/logo.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/crane_rajkumar.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/assets/logo.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/assets/logo.jpg",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/images/logo.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/images/logo.jpg",
    ]
    for url in candidates:
        try:
            if url is None:
                continue
            if os.path.exists(url):
                return url
            r = requests.get(url, timeout=5)
            if r.status_code == 200 and r.headers.get("content-type", "").startswith("image"):
                return url
        except Exception:
            continue
    return None

_logo_url = resolve_logo_url()
if _logo_url:
    try:
        st.logo(_logo_url, size="large")
    except Exception:
        st.sidebar.image(_logo_url, use_container_width=True)

# Header layout for enhanced appearance
header_cols = st.columns([1, 6])
with header_cols[0]:
    if _logo_url and os.path.exists(_logo_url):
        st.image(_logo_url, use_container_width=True)
    elif _logo_url:
        st.image(_logo_url, width=64)
with header_cols[1]:
    st.markdown("""
    <div style="display:flex; align-items:center; gap:12px;">
      <div>
        <h2 style="margin:0;">Bill Generator</h2>
        <p style="margin:0; color:#666;">A4 documents with professional layout</p>
      </div>
    </div>
    """, unsafe_allow_html=True)

# Debug logging helpers
DEBUG_VERBOSE = os.getenv("BILL_VERBOSE", "0") == "1"

def _show_celebration_once():
    if "_celebrated" not in st.session_state:
        st.session_state["_celebrated"] = True
        try:
            st.balloons()
            st.snow()
        except Exception:
            pass

def _render_landing():
    # Minimal CSS to enhance appearance
    st.markdown(
        """
        <style>
        .hero {
          padding: 28px 24px; border-radius: 14px;
          background: linear-gradient(135deg, #f0f7ff 0%, #ffffff 70%);
          border: 1px solid #e6eef8;
          box-shadow: 0 6px 24px rgba(0,0,0,0.06);
        }
        .hero h1 { margin: 0 0 8px 0; font-size: 28px; }
        .hero p { margin: 0; color: #4a5568; }
        .cta { margin-top: 16px; color: #1f6feb; }
        .tip { font-size: 13px; color: #6b7280; margin-top: 6px; }
        </style>
        """,
        unsafe_allow_html=True,
    )
    hero_cols = st.columns([3, 4])
    with hero_cols[0]:
        st.markdown(
            """
            <div class="hero">
              <h1>Generate professional A4 Bills</h1>
              <p>Uniform 10 mm margins • HTML ↔ PDF ↔ DOCX consistency • LaTeX PDFs</p>
              <div class="tip">Upload your Excel to begin. Configure premium, download merged outputs.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with hero_cols[1]:
        banner_path = os.path.join(os.getcwd(), "landing_page.png")
        if os.path.exists(banner_path):
            st.image(banner_path, use_container_width=True)
        elif _logo_url and os.path.exists(_logo_url):
            st.image(_logo_url, use_container_width=True)
    _show_celebration_once()

def _log_debug(message: str) -> None:
    if DEBUG_VERBOSE:
        try:
            st.write(message)
        except Exception:
            pass

def _log_warn(message: str) -> None:
    if DEBUG_VERBOSE:
        try:
            st.warning(message)
        except Exception:
            pass

def _log_traceback() -> None:
    if DEBUG_VERBOSE:
        try:
            st.write(traceback.format_exc())
        except Exception:
            pass

# Ensure templates directory exists
templates_dir = os.path.join(os.getcwd(), "templates")
os.makedirs(templates_dir, exist_ok=True)

# Set up Jinja2 environment
env = Environment(loader=FileSystemLoader(templates_dir), cache_size=0)

# Temporary directory
TEMP_DIR = None

def get_temp_dir():
    """Get or create a temporary directory for this session."""
    global TEMP_DIR
    if TEMP_DIR is None or not os.path.exists(TEMP_DIR):
        TEMP_DIR = tempfile.mkdtemp()
    return TEMP_DIR

def ensure_wkhtmltopdf() -> str | None:
    """Ensure wkhtmltopdf is available. If missing, download a static linux tarball and extract to a temp dir. Returns path or None."""
    # 1) Already in PATH?
    path = shutil.which("wkhtmltopdf")
    if path:
        return path
    # 2) Environment override
    env_path = os.getenv("WKHTMLTOPDF_PATH")
    if env_path and os.path.exists(env_path):
        return env_path
    # 3) Try to download a static binary (linux generic amd64)
    # Known packaging release URL (0.12.6-1) with patched Qt
    urls = [
        "https://github.com/wkhtmltopdf/packaging/releases/download/0.12.6-1/wkhtmltox-0.12.6-1_linux-generic-amd64.tar.xz",
        "https://github.com/wkhtmltopdf/packaging/releases/download/0.12.6-1/wkhtmltox-0.12.6-1.centos8.x86_64.rpm.tar.xz"
    ]
    cache_dir = os.path.join(tempfile.gettempdir(), "wkhtmltopdf_bin")
    os.makedirs(cache_dir, exist_ok=True)
    for url in urls:
        try:
            resp = requests.get(url, timeout=30, stream=True)
            if resp.status_code != 200 or len(resp.content) < 1024:
                continue
            # Validate content size to prevent memory exhaustion
            content_length = resp.headers.get('content-length')
            if content_length and int(content_length) > 100 * 1024 * 1024:  # 100MB limit
                _log_warn(f"Download too large: {content_length} bytes")
                continue
            tar_bytes = resp.content
            # Extract from .tar.xz
            with tarfile.open(fileobj=io.BytesIO(tar_bytes), mode="r:xz") as tf:  # type: ignore
                members = tf.getmembers()
                # Look for wkhtmltox/bin/wkhtmltopdf
                target_member = None
                for m in members:
                    if m.name.endswith("/wkhtmltopdf") and "/bin/" in m.name:
                        target_member = m
                        break
                if not target_member:
                    continue
                # Security check: prevent path traversal
                if ".." in target_member.name or target_member.name.startswith("/"):
                    _log_warn(f"Suspicious path in archive: {target_member.name}")
                    continue
                tf.extract(target_member, path=cache_dir)
                extracted_path = os.path.join(cache_dir, target_member.name)
                # Make executable
                os.chmod(extracted_path, 0o755)
                return extracted_path
        except Exception:
            continue
    return None

# Configure wkhtmltopdf
wkhtmltopdf_exe = None
if platform.system() == "Windows":
    wkhtmltopdf_exe = r"C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
else:
    wkhtmltopdf_exe = ensure_wkhtmltopdf() or shutil.which("wkhtmltopdf")

try:
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_exe) if wkhtmltopdf_exe else pdfkit.configuration()
except Exception:
    config = None

def number_to_words(number):
    try:
        return num2words(int(number), lang="en_IN").title()
    except:
        return str(number)

def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type, excel_data=None):
    _log_debug("Starting process_bill")
    first_page_data = {"header": [], "items": [], "totals": {}}
    last_page_data = {"payable_amount": 0, "amount_words": ""}
    deviation_data = {"items": [], "summary": {}}
    extra_items_data = {"items": []}
    note_sheet_data = {"notes": []}

    from datetime import datetime, date

    # Header (A1:G19) only — matching actual data range
    header_data = ws_wo.iloc[:19, :7].replace(np.nan, "").values.tolist()

    # Ensure all dates are formatted as date-only strings
    for i in range(len(header_data)):
        for j in range(len(header_data[i])):
            val = header_data[i][j]
            if isinstance(val, (pd.Timestamp, datetime, date)):
                header_data[i][j] = val.strftime("%d-%m-%Y")

    # Assign to first page
    first_page_data["header"] = header_data

    # Work Order items
    last_row_wo = ws_wo.shape[0]
    for i in range(21, last_row_wo):
        # Skip rows with empty description or both empty qty and rate
        if pd.isnull(ws_wo.iloc[i, 1]) or (pd.isnull(ws_wo.iloc[i, 3]) and pd.isnull(ws_wo.iloc[i, 4])):
            continue
            
        qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else 0
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                _log_warn(f"Skipping invalid quantity at Bill Quantity row {i+1}: '{qty_raw}'")
                qty = 0

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                _log_warn(f"Skipping invalid rate at Work Order row {i+1}: '{rate_raw}'")
                rate = 0

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)

    # Extra Items divider
    first_page_data["items"].append({
        "description": "Extra Items (With Premium)",
        "bold": True,
        "underline": True,
        "amount": 0,
        "quantity": 0,
        "rate": 0,
        "serial_no": "",
        "unit": "",
        "remark": "",
        "is_divider": True
    })

    # Extra Items  
    last_row_extra = ws_extra.shape[0]
    for j in range(5, last_row_extra):
        # Skip rows with empty description or serial number
        if pd.isnull(ws_extra.iloc[j, 2]) or pd.isnull(ws_extra.iloc[j, 0]):
            continue
            
        qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else 0
        rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                _log_warn(f"Skipping invalid quantity at Extra Items row {j+1}: '{qty_raw}'")
                qty = 0

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                _log_warn(f"Skipping invalid rate at Extra Items row {j+1}: '{rate_raw}'")
                rate = 0

        # Apply premium to extra item rate  
        if premium_type == "Percentage":
            rate_with_premium = rate * (1 + premium_percent / 100)
        else:
            rate_with_premium = rate + premium_percent

        item = {
            "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
            "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
            "unit": str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else "",
            "quantity": qty,
            "rate": rate_with_premium,
            "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
            "amount": round(qty * rate_with_premium) if qty and rate_with_premium else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)
        extra_items_data["items"].append(item.copy())

    # Totals - Calculate work order and extra items separately
    work_order_items = []
    extra_items = []
    divider_found = False
    
    for item in first_page_data["items"]:
        if item.get("is_divider") and "Extra Items" in item.get("description", ""):
            divider_found = True
            continue
        if not item.get("is_divider"):
            if divider_found:
                extra_items.append(item)
            else:
                work_order_items.append(item)
    
    total_work_order = sum(item.get("amount", 0) for item in work_order_items)
    total_extra_items = sum(item.get("amount", 0) for item in extra_items)
    grand_total = total_work_order + total_extra_items

    first_page_data["totals"] = {
        "total_work_order": total_work_order,
        "total_extra_items": total_extra_items,
        "grand_total": grand_total
    }

    last_page_data = {
        "payable_amount": grand_total,
        "amount_words": number_to_words(grand_total)
    }

    # Deviation Statement - exactly as per original code
    work_order_total = 0
    executed_total = 0
    overall_excess = 0
    overall_saving = 0
    for i in range(21, last_row_wo):
        qty_wo_raw = ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else 0
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None
        qty_bill_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else 0

        qty_wo = 0
        if isinstance(qty_wo_raw, (int, float)):
            qty_wo = float(qty_wo_raw)
        elif isinstance(qty_wo_raw, str):
            cleaned_qty_wo = qty_wo_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_wo = float(cleaned_qty_wo)
            except ValueError:
                qty_wo = 0

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                rate = 0

        qty_bill = 0
        if isinstance(qty_bill_raw, (int, float)):
            qty_bill = float(qty_bill_raw)
        elif isinstance(qty_bill_raw, str):
            cleaned_qty_bill = qty_bill_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_bill = float(cleaned_qty_bill)
            except ValueError:
                qty_bill = 0

        amt_wo = round(qty_wo * rate)
        amt_bill = round(qty_bill * rate)
        excess_qty = qty_bill - qty_wo if qty_bill > qty_wo else 0
        excess_amt = round(excess_qty * rate) if excess_qty > 0 else 0
        saving_qty = qty_wo - qty_bill if qty_bill < qty_wo else 0
        saving_amt = round(saving_qty * rate) if saving_qty > 0 else 0

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "qty_wo": qty_wo,
            "rate": rate,
            "amt_wo": amt_wo,
            "qty_bill": qty_bill,
            "amt_bill": amt_bill,
            "excess_qty": excess_qty,
            "excess_amt": excess_amt,
            "saving_qty": saving_qty,
            "saving_amt": saving_amt
        }
        deviation_data["items"].append(item)
        work_order_total += amt_wo
        executed_total += amt_bill
        overall_excess += excess_amt
        overall_saving += saving_amt

    # Deviation Summary
    tender_premium_f = round(work_order_total * (premium_percent / 100) if premium_type == "above" else -work_order_total * (premium_percent / 100))
    tender_premium_h = round(executed_total * (premium_percent / 100) if premium_type == "above" else -executed_total * (premium_percent / 100))
    tender_premium_j = round(overall_excess * (premium_percent / 100) if premium_type == "above" else -overall_excess * (premium_percent / 100))
    tender_premium_l = round(overall_saving * (premium_percent / 100) if premium_type == "above" else -overall_saving * (premium_percent / 100))
    grand_total_f = work_order_total + tender_premium_f
    grand_total_h = executed_total + tender_premium_h
    grand_total_j = overall_excess + tender_premium_j
    grand_total_l = overall_saving + tender_premium_l
    net_difference = grand_total_h - grand_total_f

    deviation_data["summary"] = {
        "work_order_total": round(work_order_total),
        "executed_total": round(executed_total),
        "overall_excess": round(overall_excess),
        "overall_saving": round(overall_saving),
        "premium": {"percent": premium_percent / 100, "type": premium_type},
        "tender_premium_f": tender_premium_f,
        "tender_premium_h": tender_premium_h,
        "tender_premium_j": tender_premium_j,
        "tender_premium_l": tender_premium_l,
        "grand_total_f": grand_total_f,
        "grand_total_h": grand_total_h,
        "grand_total_j": grand_total_j,
        "grand_total_l": grand_total_l,
        "net_difference": round(net_difference)
    }

    # Note sheet data from Excel NOTE SHEET if available
    note_sheet_data["notes"] = []
    
    # Try to read from NOTE SHEET in Excel first
    try:
        ws_notes = excel_data.get('NOTE SHEET', pd.DataFrame()) if excel_data is not None else pd.DataFrame()
        if not ws_notes.empty:
            # Extract structured data from NOTE SHEET
            for i in range(2, min(ws_notes.shape[0], 25)):  # Start from row 2, limit to reasonable range
                description = str(ws_notes.iloc[i, 1]) if pd.notnull(ws_notes.iloc[i, 1]) else ""
                details = str(ws_notes.iloc[i, 2]) if pd.notnull(ws_notes.iloc[i, 2]) else ""
                
                if description and description != "nan" and len(description.strip()) > 2:
                    if details and details != "nan" and len(details.strip()) > 0:
                        combined_note = f"{description.strip()}: {details.strip()}"
                    else:
                        combined_note = description.strip()
                    
                    if combined_note not in note_sheet_data["notes"]:
                        note_sheet_data["notes"].append(combined_note)
    except Exception:
        pass
    
    # If no notes found in Excel, provide basic ones
    if not note_sheet_data["notes"]:
        note_sheet_data["notes"] = [
            "All work executed as per approved drawings and specifications.",
            "Quality control measures implemented throughout execution.", 
            "Safety protocols followed as per government standards."
        ]

    return first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data

def generate_html_from_template(template_name, context):
    """Generate HTML content from template with proper context"""
    try:
        # Create fresh environment to avoid template caching issues
        fresh_env = Environment(loader=FileSystemLoader(templates_dir), cache_size=0)
        template = fresh_env.get_template(template_name)
        html_content = template.render(**context)
        return html_content
    except Exception as e:
        _log_warn(f"Template rendering error for {template_name}: {str(e)}")
        return f"<html><body><h1>Template Error</h1><p>{str(e)}</p></body></html>"

def generate_pdf_from_html(html_content, output_path, landscape=False):
    """Generate PDF from HTML content"""
    try:
        if config:
            options = {
                'page-size': 'A4',
                'orientation': 'Landscape' if landscape else 'Portrait',
                'margin-top': '10mm',
                'margin-right': '10mm',
                'margin-bottom': '10mm',
                'margin-left': '10mm',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None,
                'print-media-type': None,
                'disable-smart-shrinking': None,
                'zoom': '1.0',
                'dpi': 96,
                'minimum-font-size': 12
            }
            pdfkit.from_string(html_content, output_path, options=options, configuration=config)
            return True
        else:
            # Fallback to xhtml2pdf with proper scaling
            # Adjust layout for xhtml2pdf to avoid narrow content
            fallback_html = html_content
            for mm in ("190mm", "277mm"):
                fallback_html = fallback_html.replace(f"width: {mm}", "width: 100%")
                fallback_html = fallback_html.replace(f"max-width: {mm}", "width: 100%")
            # Remove centered margin that can introduce side gaps
            fallback_html = fallback_html.replace("margin: 0 auto;", "margin: 0;")
            orientation = 'landscape' if landscape else 'portrait'
            default_css = f'@page {{ size: A4 {orientation}; margin: 10mm; }}'
            with open(output_path, "w+b") as result_file:
                pisa_status = pisa.CreatePDF(fallback_html.encode('utf-8'), dest=result_file, default_css=default_css)
                return pisa_status.err == 0
    except Exception as e:
        _log_warn(f"PDF generation error: {str(e)}")
        return False

def generate_docx_from_data(data, output_path):
    """Generate DOCX from data"""
    try:
        doc = Document()
        
        # Set margins
        section = doc.sections[0]
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        
        # Add header information
        if "header" in data:
            for row in data["header"]:
                if any(cell for cell in row if str(cell).strip()):
                    p = doc.add_paragraph()
                    p.add_run(" | ".join(str(cell) for cell in row if str(cell).strip()))
        
        # Add items table
        if "items" in data and data["items"]:
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'S.No.'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Unit'
            hdr_cells[3].text = 'Quantity'
            hdr_cells[4].text = 'Rate'
            hdr_cells[5].text = 'Amount'
            hdr_cells[6].text = 'Remark'
            
            # Data rows
            for item in data["items"]:
                if item.get("is_divider"):
                    row_cells = table.add_row().cells
                    row_cells[1].text = item.get("description", "")
                    # Make divider bold
                    for paragraph in row_cells[1].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get("serial_no", ""))
                    row_cells[1].text = str(item.get("description", ""))
                    row_cells[2].text = str(item.get("unit", ""))
                    row_cells[3].text = str(item.get("quantity", ""))
                    row_cells[4].text = str(item.get("rate", ""))
                    row_cells[5].text = str(item.get("amount", ""))
                    row_cells[6].text = str(item.get("remark", ""))
        
        # Add totals
        if "totals" in data:
            doc.add_paragraph(f"Total Work Order: {data['totals'].get('total_work_order', 0)}")
            doc.add_paragraph(f"Total Extra Items: {data['totals'].get('total_extra_items', 0)}")
            doc.add_paragraph(f"Grand Total: {data['totals'].get('grand_total', 0)}")
        
        doc.save(output_path)
        return True
    except Exception as e:
        _log_warn(f"DOCX generation error: {str(e)}")
        return False

# Main Streamlit app
def main():
    if "uploaded_file" not in st.session_state:
        _render_landing()
        
        st.markdown("### Upload Excel File")
        uploaded_file = st.file_uploader("Choose an Excel file", type=['xlsx', 'xls'])
        
        if uploaded_file:
            st.session_state.uploaded_file = uploaded_file
            st.rerun()
        return

    uploaded_file = st.session_state.uploaded_file
    
    # Premium configuration
    col1, col2 = st.columns(2)
    with col1:
        premium_percent = st.number_input("Premium Value (%)", min_value=-99.99, max_value=99.99, value=10.0, step=0.01)
    with col2:
        premium_type = st.selectbox("Premium Type", ["Percentage", "Fixed Amount"])
        st.caption("Premium range: -99.99% to +99.99% (negative = below, positive = above)")

    if st.button("Generate Bills"):
        try:
            # Read Excel file
            excel_data = pd.read_excel(uploaded_file, sheet_name=None)
            
            # Extract worksheets - use exact names from Excel file
            ws_wo = excel_data.get('Work Order', pd.DataFrame())
            ws_bq = excel_data.get('Bill Quantity', pd.DataFrame())  
            ws_extra = excel_data.get('Extra Items', pd.DataFrame())
            
            if ws_wo.empty or ws_bq.empty or ws_extra.empty:
                st.error("Required worksheets not found. Please ensure Excel file contains 'Work Order', 'Bill Quantity', and 'Extra Items' sheets.")
                return
            
            # Process the bill data
            first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data = process_bill(
                ws_wo, ws_bq, ws_extra, premium_percent, premium_type, excel_data
            )
            
            # Generate documents
            temp_dir = get_temp_dir()
            
            # HTML generation with correct variable mapping
            first_page_html = generate_html_from_template("first_page.html", {
                "header": first_page_data["header"],
                "items": first_page_data["items"],
                "totals": first_page_data["totals"]
            })
            last_page_html = generate_html_from_template("last_page.html", {
                "payable_amount": last_page_data["payable_amount"],
                "amount_words": last_page_data["amount_words"]
            })
            deviation_html = generate_html_from_template("deviation_sheet.html", {
                "items": deviation_data["items"],
                "summary": deviation_data["summary"]
            })
            extra_items_html = generate_html_from_template("extra_items_sheet.html", {
                "items": extra_items_data["items"]
            })
            note_sheet_html = generate_html_from_template("note_sheet.html", {
                "notes": note_sheet_data["notes"]
            })
            
            # PDF generation with proper orientation
            pdf_files = []
            pdf_configs = [
                ("first_page", first_page_html, False),
                ("last_page", last_page_html, False),
                ("deviation_sheet", deviation_html, True),  # Landscape
                ("extra_items_sheet", extra_items_html, False),
                ("note_sheet", note_sheet_html, False)
            ]
            
            for name, html_content, landscape in pdf_configs:
                pdf_path = os.path.join(temp_dir, f"{name}.pdf")
                st.write(f"Generating {name}...")
                success = generate_pdf_from_html(html_content, pdf_path, landscape)
                if success:
                    pdf_files.append(pdf_path)
                    st.write(f"✓ {name} generated successfully")
                else:
                    st.write(f"✗ {name} failed to generate")
            
            # DOCX generation
            docx_path = os.path.join(temp_dir, "bill_complete.docx")
            generate_docx_from_data(first_page_data, docx_path)
            
            # Display results and debug info
            st.success(f"Bills generated successfully! Generated {len(pdf_files)} PDF files.")
            st.write(f"PDF files created: {[os.path.basename(f) for f in pdf_files]}")
            
            # Download buttons
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if pdf_files:
                    # Create merged PDF
                    merged_pdf_path = os.path.join(temp_dir, "merged_bill.pdf")
                    writer = PdfWriter()
                    for pdf_file in pdf_files:
                        if os.path.exists(pdf_file):
                            reader = PdfReader(pdf_file)
                            for page in reader.pages:
                                writer.add_page(page)
                    
                    with open(merged_pdf_path, "wb") as output_file:
                        writer.write(output_file)
                    
                    with open(merged_pdf_path, "rb") as file:
                        st.download_button("Download PDF", file.read(), "bill.pdf", "application/pdf")
            
            with col2:
                if os.path.exists(docx_path):
                    with open(docx_path, "rb") as file:
                        st.download_button("Download DOCX", file.read(), "bill.docx", 
                                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            with col3:
                # Create ZIP with all files
                zip_path = os.path.join(temp_dir, "all_bills.zip")
                with zipfile.ZipFile(zip_path, 'w') as zipf:
                    for pdf_file in pdf_files:
                        if os.path.exists(pdf_file):
                            zipf.write(pdf_file, os.path.basename(pdf_file))
                    if os.path.exists(docx_path):
                        zipf.write(docx_path, os.path.basename(docx_path))
                
                with open(zip_path, "rb") as file:
                    st.download_button("Download All", file.read(), "all_bills.zip", "application/zip")
            
            # Preview all generated content
            st.markdown("### Preview")
            
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["First Page", "Last Page", "Deviation", "Extra Items", "Note Sheet"])
            
            with tab1:
                st.components.v1.html(first_page_html, height=600, scrolling=True)
                
            with tab2:
                st.components.v1.html(last_page_html, height=600, scrolling=True)
                
            with tab3:
                st.components.v1.html(deviation_html, height=600, scrolling=True)
                
            with tab4:
                st.components.v1.html(extra_items_html, height=600, scrolling=True)
                
            with tab5:
                st.components.v1.html(note_sheet_html, height=600, scrolling=True)
            
        except Exception as e:
            st.error(f"Error generating bills: {str(e)}")
            _log_traceback()

if __name__ == "__main__":
    main()
