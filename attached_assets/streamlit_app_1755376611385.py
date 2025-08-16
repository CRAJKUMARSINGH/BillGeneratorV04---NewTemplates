import streamlit as st
import pandas as pd
import pdfkit
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Mm
from num2words import num2words
import os
import zipfile
import tempfile
from jinja2 import Environment, FileSystemLoader
from pypdf import PdfReader, PdfWriter
import numpy as np
import platform
from datetime import datetime
import subprocess
import shutil
import traceback
import argparse
import requests
from xhtml2pdf import pisa
import tarfile
import io

# Page setup and branding
_local_logo = os.path.join(os.getcwd(), "crane_rajkumar.png")
_page_icon = _local_logo if os.path.exists(_local_logo) else "📄"
st.set_page_config(page_title="Bill Generator", page_icon=_page_icon, layout="wide")

def resolve_logo_url() -> str | None:
    candidates = [
        # Prefer local crane_rajkumar.png if present
        _local_logo if os.path.exists(_local_logo) else None,
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/logo.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/crane_rajkumar.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/assets/logo.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/assets/logo.jpg",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/images/logo.png",
        "https://raw.githubusercontent.com/CRAJKUMARSINGH/Priyanka_TenderV01/HEAD/images/logo.jpg",
    ]
    for url in candidates:
        try:
            if url is None:
                continue
            if os.path.exists(url):
                return url
            r = requests.get(url, timeout=5)
            if r.status_code == 200 and r.headers.get("content-type", "").startswith("image"):
                return url
        except Exception:
            continue
    return None

_logo_url = resolve_logo_url()
if _logo_url:
    try:
        st.logo(_logo_url, size="large")
    except Exception:
        st.sidebar.image(_logo_url, use_container_width=True)

# Header layout for enhanced appearance
header_cols = st.columns([1, 6])
with header_cols[0]:
    if _logo_url and os.path.exists(_logo_url):
        st.image(_logo_url, use_container_width=True)
    elif _logo_url:
        st.image(_logo_url, width=64)
with header_cols[1]:
    st.markdown("""
    <div style="display:flex; align-items:center; gap:12px;">
      <div>
        <h2 style="margin:0;">Bill Generator</h2>
        <p style="margin:0; color:#666;">A4 documents with professional layout</p>
      </div>
    </div>
    """, unsafe_allow_html=True)

# Debug logging helpers
DEBUG_VERBOSE = os.getenv("BILL_VERBOSE", "0") == "1"

def _show_celebration_once():
    if "_celebrated" not in st.session_state:
        st.session_state["_celebrated"] = True
        try:
            st.balloons()
            st.snow()
        except Exception:
            pass

def _render_landing():
    # Minimal CSS to enhance appearance
    st.markdown(
        """
        <style>
        .hero {
          padding: 28px 24px; border-radius: 14px;
          background: linear-gradient(135deg, #f0f7ff 0%, #ffffff 70%);
          border: 1px solid #e6eef8;
          box-shadow: 0 6px 24px rgba(0,0,0,0.06);
        }
        .hero h1 { margin: 0 0 8px 0; font-size: 28px; }
        .hero p { margin: 0; color: #4a5568; }
        .cta { margin-top: 16px; color: #1f6feb; }
        .tip { font-size: 13px; color: #6b7280; margin-top: 6px; }
        </style>
        """,
        unsafe_allow_html=True,
    )
    hero_cols = st.columns([3, 4])
    with hero_cols[0]:
        st.markdown(
            """
            <div class="hero">
              <h1>Generate professional A4 Bills</h1>
              <p>Uniform 10 mm margins • HTML ↔ PDF ↔ DOCX consistency • LaTeX PDFs</p>
              <div class="tip">Upload your Excel to begin. Configure premium, download merged outputs.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with hero_cols[1]:
        banner_path = os.path.join(os.getcwd(), "landing_page.png")
        if os.path.exists(banner_path):
            st.image(banner_path, use_container_width=True)
        elif _logo_url and os.path.exists(_logo_url):
            st.image(_logo_url, use_container_width=True)
    _show_celebration_once()

def _log_debug(message: str) -> None:
    if DEBUG_VERBOSE:
        try:
            st.write(message)
        except Exception:
            pass

def _log_warn(message: str) -> None:
    if DEBUG_VERBOSE:
        try:
            st.warning(message)
        except Exception:
            pass

def _log_traceback() -> None:
    if DEBUG_VERBOSE:
        try:
            st.write(traceback.format_exc())
        except Exception:
            pass

# Set up Jinja2 environment
env = Environment(loader=FileSystemLoader("templates"), cache_size=0)

# Temporary directory
TEMP_DIR = None

def get_temp_dir():
    """Get or create a temporary directory for this session."""
    global TEMP_DIR
    if TEMP_DIR is None or not os.path.exists(TEMP_DIR):
        TEMP_DIR = tempfile.mkdtemp()
    return TEMP_DIR
def ensure_wkhtmltopdf() -> str | None:
    """Ensure wkhtmltopdf is available. If missing, download a static linux tarball and extract to a temp dir. Returns path or None."""
    # 1) Already in PATH?
    path = shutil.which("wkhtmltopdf")
    if path:
        return path
    # 2) Environment override
    env_path = os.getenv("WKHTMLTOPDF_PATH")
    if env_path and os.path.exists(env_path):
        return env_path
    # 3) Try to download a static binary (linux generic amd64)
    # Known packaging release URL (0.12.6-1) with patched Qt
    urls = [
        "https://github.com/wkhtmltopdf/packaging/releases/download/0.12.6-1/wkhtmltox-0.12.6-1_linux-generic-amd64.tar.xz",
        "https://github.com/wkhtmltopdf/packaging/releases/download/0.12.6-1/wkhtmltox-0.12.6-1.centos8.x86_64.rpm.tar.xz"
    ]
    cache_dir = os.path.join(tempfile.gettempdir(), "wkhtmltopdf_bin")
    os.makedirs(cache_dir, exist_ok=True)
    for url in urls:
        try:
            resp = requests.get(url, timeout=30, stream=True)
            if resp.status_code != 200 or len(resp.content) < 1024:
                continue
            # Validate content size to prevent memory exhaustion
            content_length = resp.headers.get('content-length')
            if content_length and int(content_length) > 100 * 1024 * 1024:  # 100MB limit
                _log_warn(f"Download too large: {content_length} bytes")
                continue
            tar_bytes = resp.content
            # Extract from .tar.xz
            with tarfile.open(fileobj=io.BytesIO(tar_bytes), mode="r:xz") as tf:  # type: ignore
                members = tf.getmembers()
                # Look for wkhtmltox/bin/wkhtmltopdf
                target_member = None
                for m in members:
                    if m.name.endswith("/wkhtmltopdf") and "/bin/" in m.name:
                        target_member = m
                        break
                if not target_member:
                    continue
                # Security check: prevent path traversal
                if ".." in target_member.name or target_member.name.startswith("/"):
                    _log_warn(f"Suspicious path in archive: {target_member.name}")
                    continue
                tf.extract(target_member, path=cache_dir)
                extracted_path = os.path.join(cache_dir, target_member.name)
                # Make executable
                os.chmod(extracted_path, 0o755)
                return extracted_path
        except Exception:
            continue
    return None

# Configure wkhtmltopdf
wkhtmltopdf_exe = None
if platform.system() == "Windows":
    wkhtmltopdf_exe = r"C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
else:
    wkhtmltopdf_exe = ensure_wkhtmltopdf() or shutil.which("wkhtmltopdf")

try:
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_exe) if wkhtmltopdf_exe else pdfkit.configuration()
except Exception:
    config = None

def number_to_words(number):
    try:
        return num2words(int(number), lang="en_IN").title()
    except:
        return str(number)
##########################################################################################
def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type):
    _log_debug("Starting process_bill")
    first_page_data = {"header": [], "items": [], "totals": {}}
    last_page_data = {"payable_amount": 0, "amount_words": ""}
    deviation_data = {"items": [], "summary": {}}
    extra_items_data = {"items": []}
    note_sheet_data = {"notes": []}
################################################################################################################
    # Header (A1:I19)
    #header_data = ws_wo.iloc[:19].replace(np.nan, "").values.tolist()
    #first_page_data["header"] = header_data
    ###### REPLACEMENT 18 APRIL 2025
    from datetime import datetime, date

    # Header (A1:G19) only — matching actual data range
    header_data = ws_wo.iloc[:19, :7].replace(np.nan, "").values.tolist()

    # Ensure all dates are formatted as date-only strings (optional step, if needed before saving)
    for i in range(len(header_data)):
        for j in range(len(header_data[i])):
            val = header_data[i][j]
            if isinstance(val, (pd.Timestamp, datetime, date)):
                header_data[i][j] = val.strftime("%d-%m-%Y")

    # Assign to first page
    first_page_data["header"] = header_data
############################################################################################################
    # Work Order items
    last_row_wo = ws_wo.shape[0]
    for i in range(21, last_row_wo):
        qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else 0
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                _log_warn(f"Skipping invalid quantity at Bill Quantity row {i+1}: '{qty_raw}'")
                qty = 0

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                _log_warn(f"Skipping invalid rate at Work Order row {i+1}: '{rate_raw}'")
                rate = 0

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)

    # Extra Items divider
    first_page_data["items"].append({
        "description": "Extra Items (With Premium)",
        "bold": True,
        "underline": True,
        "amount": 0,
        "quantity": 0,
        "rate": 0,
        "serial_no": "",
        "unit": "",
        "remark": "",
        "is_divider": True
    })

    # Extra Items
    last_row_extra = ws_extra.shape[0]
    for j in range(6, last_row_extra):
        qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else 0
        rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

        qty = 0
        if isinstance(qty_raw, (int, float)):
            qty = float(qty_raw)
        elif isinstance(qty_raw, str):
            cleaned_qty = qty_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty = float(cleaned_qty)
            except ValueError:
                _log_warn(f"Skipping invalid quantity at Extra Items row {j+1}: '{qty_raw}'")
                qty = 0

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                _log_warn(f"Skipping invalid rate at Extra Items row {j+1}: '{rate_raw}'")
                rate = 0

        item = {
            "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
            "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
            "unit": str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else "",
            "quantity": qty,
            "rate": rate,
            "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
            "amount": round(qty * rate) if qty and rate else 0,
            "is_divider": False
        }
        first_page_data["items"].append(item)
        extra_items_data["items"].append(item.copy())  # Copy for standalone Extra Items

    # Totals
    data_items = [item for item in first_page_data["items"] if not item.get("is_divider", False)]
    total_amount = round(sum(item.get("amount", 0) for item in data_items))
    premium_amount = round(total_amount * (premium_percent / 100) if premium_type == "above" else -total_amount * (premium_percent / 100))
    payable_amount = round(total_amount + premium_amount)

    first_page_data["totals"] = {
        "grand_total": total_amount,
        "premium": {"percent": premium_percent / 100, "type": premium_type, "amount": premium_amount},
        "payable": payable_amount
    }

    try:
        extra_items_start = next(i for i, item in enumerate(first_page_data["items"]) if item.get("description") == "Extra Items (With Premium)")
        extra_items = [item for item in first_page_data["items"][extra_items_start + 1:] if not item.get("is_divider", False)]
        extra_items_sum = round(sum(item.get("amount", 0) for item in extra_items))
        extra_items_premium = round(extra_items_sum * (premium_percent / 100) if premium_type == "above" else -extra_items_sum * (premium_percent / 100))
        first_page_data["totals"]["extra_items_sum"] = extra_items_sum + extra_items_premium
    except StopIteration:
        first_page_data["totals"]["extra_items_sum"] = 0

    # Last Page
    last_page_data = {"payable_amount": payable_amount, "amount_words": number_to_words(payable_amount)}

    # Deviation Statement
    work_order_total = 0
    executed_total = 0
    overall_excess = 0
    overall_saving = 0
    for i in range(21, last_row_wo):
        _log_debug(f"Processing deviation row {i+1}: wo_qty={ws_wo.iloc[i, 3]}, wo_rate={ws_wo.iloc[i, 4]}, bq_qty={ws_bq.iloc[i, 3] if i < ws_bq.shape[0] else 'N/A'}")
        qty_wo_raw = ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else 0
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None
        qty_bill_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else 0

        qty_wo = 0
        if isinstance(qty_wo_raw, (int, float)):
            qty_wo = float(qty_wo_raw)
        elif isinstance(qty_wo_raw, str):
            cleaned_qty_wo = qty_wo_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_wo = float(cleaned_qty_wo)
            except ValueError:
                _log_warn(f"Skipping invalid qty_wo at row {i+1}: '{qty_wo_raw}'")
                qty_wo = 0

        rate = 0
        if isinstance(rate_raw, (int, float)):
            rate = float(rate_raw)
        elif isinstance(rate_raw, str):
            cleaned_rate = rate_raw.strip().replace(',', '').replace(' ', '')
            try:
                rate = float(cleaned_rate)
            except ValueError:
                _log_warn(f"Skipping invalid rate at row {i+1}: '{rate_raw}'")
                rate = 0

        qty_bill = 0
        if isinstance(qty_bill_raw, (int, float)):
            qty_bill = float(qty_bill_raw)
        elif isinstance(qty_bill_raw, str):
            cleaned_qty_bill = qty_bill_raw.strip().replace(',', '').replace(' ', '')
            try:
                qty_bill = float(cleaned_qty_bill)
            except ValueError:
                _log_warn(f"Skipping invalid qty_bill at row {i+1}: '{qty_bill_raw}'")
                qty_bill = 0

        amt_wo = round(qty_wo * rate)
        amt_bill = round(qty_bill * rate)
        excess_qty = qty_bill - qty_wo if qty_bill > qty_wo else 0
        excess_amt = round(excess_qty * rate) if excess_qty > 0 else 0
        saving_qty = qty_wo - qty_bill if qty_bill < qty_wo else 0
        saving_amt = round(saving_qty * rate) if saving_qty > 0 else 0

        item = {
            "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
            "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
            "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
            "qty_wo": qty_wo,
            "rate": rate,
            "amt_wo": amt_wo,
            "qty_bill": qty_bill,
            "amt_bill": amt_bill,
            "excess_qty": excess_qty,
            "excess_amt": excess_amt,
            "saving_qty": saving_qty,
            "saving_amt": saving_amt
        }
        deviation_data["items"].append(item)
        work_order_total += amt_wo
        executed_total += amt_bill
        overall_excess += excess_amt
        overall_saving += saving_amt

    # Deviation Summary
    tender_premium_f = round(work_order_total * (premium_percent / 100) if premium_type == "above" else -work_order_total * (premium_percent / 100))
    tender_premium_h = round(executed_total * (premium_percent / 100) if premium_type == "above" else -executed_total * (premium_percent / 100))
    tender_premium_j = round(overall_excess * (premium_percent / 100) if premium_type == "above" else -overall_excess * (premium_percent / 100))
    tender_premium_l = round(overall_saving * (premium_percent / 100) if premium_type == "above" else -overall_saving * (premium_percent / 100))
    grand_total_f = work_order_total + tender_premium_f
    grand_total_h = executed_total + tender_premium_h
    grand_total_j = overall_excess + tender_premium_j
    grand_total_l = overall_saving + tender_premium_l
    net_difference = grand_total_h - grand_total_f

    deviation_data["summary"] = {
        "work_order_total": round(work_order_total),
        "executed_total": round(executed_total),
        "overall_excess": round(overall_excess),
        "overall_saving": round(overall_saving),
        "premium": {"percent": premium_percent / 100, "type": premium_type},
        "tender_premium_f": tender_premium_f,
        "tender_premium_h": tender_premium_h,
        "tender_premium_j": tender_premium_j,
        "tender_premium_l": tender_premium_l,
        "grand_total_f": grand_total_f,
        "grand_total_h": grand_total_h,
        "grand_total_j": grand_total_j,
        "grand_total_l": grand_total_l,
        "net_difference": round(net_difference)
    }

    _log_debug("Prepared first_page_data items")
    _log_debug("Prepared extra_items_data items")
    _log_debug("Prepared deviation_data items")
    return first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data
########################################################################################################################################################
def generate_bill_notes(payable_amount, work_order_amount, extra_item_amount):
    percentage_work_done = float(payable_amount / work_order_amount * 100) if work_order_amount > 0 else 0
    serial_number = 1
    note = []
    note.append(f"{serial_number}. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount.")
    serial_number += 1
    if percentage_work_done < 90:
        note.append(f"{serial_number}. The execution of work at final stage is less than 90%...")
        serial_number += 1
    elif percentage_work_done > 100 and percentage_work_done <= 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    elif percentage_work_done > 105:
        note.append(f"{serial_number}. Requisite Deviation Statement is enclosed...")
        serial_number += 1
    note.append(f"{serial_number}. Quality Control (QC) test reports attached.")
    serial_number += 1
    if extra_item_amount > 0:
        extra_item_percentage = float(extra_item_amount / work_order_amount * 100) if work_order_amount > 0 else 0
        if extra_item_percentage > 5:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        else:
            note.append(f"{serial_number}. The amount of Extra items is Rs. {extra_item_amount}...")
        serial_number += 1
    note.append(f"{serial_number}. Please peruse above details for necessary decision-making.")
    note.append("")
    note.append("                                Premlata Jain")
    note.append("                               AAO- As Auditor")
    return {"notes": note}

def generate_pdf(sheet_name, data, orientation, output_path):
    _log_debug(f"Generating PDF for {sheet_name}")
    try:
        template = env.get_template(f"{sheet_name.lower().replace(' ', '_')}.html")
        html_content = template.render(data=data)
        # Save HTML alongside PDF for consistency checks/comparison
        try:
            html_path = output_path[:-4] + ".html" if output_path.lower().endswith(".pdf") else output_path + ".html"
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
        except Exception:
            pass
        options = {
            "page-size": "A4",
            "orientation": orientation,
            "margin-top": "10mm",
            "margin-bottom": "10mm",
            "margin-left": "10mm",
            "margin-right": "10mm",
            "print-media-type": None,
            "enable-local-file-access": None,
            "disable-smart-shrinking": None,
            "zoom": "1",
            "dpi": 300,
        }
        if config is None:
            _log_warn("wkhtmltopdf missing; using xhtml2pdf fallback.")
            try:
                # Adjust layout for xhtml2pdf to avoid narrow content due to mm widths
                fallback_html = html_content
                for mm in ("190mm", "277mm"):
                    fallback_html = fallback_html.replace(f"width: {mm}", "width: 100%")
                    fallback_html = fallback_html.replace(f"max-width: {mm}", "width: 100%")
                # Remove centered margin that can introduce side gaps
                fallback_html = fallback_html.replace("margin: 0 auto;", "margin: 0;")
                default_css = '@page { size: A4 %s; margin: 10mm; }' % ("landscape" if orientation=="landscape" else "portrait")
                with open(output_path, "wb") as pdf_file:
                    pisa.CreatePDF(src=fallback_html, dest=pdf_file, default_css=default_css)
            except Exception:
                _log_traceback()
        else:
            pdfkit.from_string(
                html_content,
                output_path,
                configuration=config,
                options=options
            )
        _log_debug(f"Finished PDF for {sheet_name}")
    except Exception as e:
        st.error(f"Error generating PDF for {sheet_name}: {str(e)}")
        _log_traceback()
        raise

def compile_latex_templates(output_dir: str):
    """Compile LaTeX templates in `LaTeX_Templates` to PDFs into `output_dir`. Returns list of PDF paths."""
    _log_debug("Compiling LaTeX templates to PDF...")
    os.makedirs(output_dir, exist_ok=True)
    latex_dir = os.path.join(os.getcwd(), "LaTeX_Templates")
    if not os.path.isdir(latex_dir):
        _log_warn("LaTeX templates directory not found.")
        return []
    if shutil.which("pdflatex") is None:
        _log_warn("pdflatex not found on system. Skipping LaTeX PDF compilation.")
        return []
    compiled_pdfs = []
    for name in os.listdir(latex_dir):
        if not name.lower().endswith(".tex"):
            continue
        tex_path = os.path.join(latex_dir, name)
        try:
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-halt-on-error", f"-output-directory={output_dir}", tex_path],
                cwd=latex_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                check=True,
                text=True,
            )
            pdf_name = os.path.splitext(name)[0] + ".pdf"
            pdf_path = os.path.join(output_dir, pdf_name)
            if os.path.exists(pdf_path):
                compiled_pdfs.append(pdf_path)
        except Exception as e:
            _log_warn(f"Failed to compile {name}: {e}")
    return compiled_pdfs

def create_word_doc(sheet_name, data, doc_path):
    _log_debug(f"Creating Word doc for {sheet_name}")
    try:
        doc = Document()
        # Page setup: A4, 10mm margins, orientation by sheet
        section = doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        # default portrait; switch later if needed
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        if sheet_name == "Deviation Statement":
            # landscape
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Mm(297)
            section.page_height = Mm(210)
        if sheet_name == "First Page": 
            # Mirror HTML table columns: Unit, Qty since last, Qty upto date, Item No., Description, Rate, Amount upto date, Amount since prev, Remark
            table = doc.add_table(rows=len(data["items"]) + 3, cols=9)
            table.style = "Table Grid"
            table.autofit = False
            # Column widths in mm to total 190mm content width
            fp_widths_mm = [12, 18, 18, 12, 80, 12, 18, 12, 8]
            for idx, w in enumerate(fp_widths_mm):
                for row in table.rows:
                    row.cells[idx].width = Mm(w)
            for i, item in enumerate(data["items"]):
                row = table.rows[i]
                row.cells[0].text = str(item.get("unit", ""))
                row.cells[1].text = str(item.get("quantity_since_last", ""))
                qty_upto = item.get("quantity_upto_date") if item.get("quantity_upto_date") not in (None, "") else item.get("quantity", "")
                row.cells[2].text = str(qty_upto)
                row.cells[3].text = str(item.get("serial_no", ""))
                row.cells[4].text = str(item.get("description", ""))
                row.cells[5].text = str(item.get("rate", ""))
                row.cells[6].text = str(item.get("amount", ""))
                row.cells[7].text = str(item.get("amount_previous", ""))
                row.cells[8].text = str(item.get("remark", ""))
            row = table.rows[-3]
            row.cells[4].text = "Grand Total"
            row.cells[6].text = str(data["totals"].get("grand_total", ""))
            row = table.rows[-2]
            row.cells[4].text = f"Tender Premium @ {data['totals']['premium'].get('percent', 0):.2%}"
            row.cells[6].text = str(data["totals"]["premium"].get("amount", ""))
            row = table.rows[-1]
            row.cells[4].text = "Payable Amount"
            row.cells[6].text = str(data["totals"].get("payable", ""))
        elif sheet_name == "Last Page":
            doc.add_paragraph(f"Payable Amount: {data.get('payable_amount', '')}")
            doc.add_paragraph(f"Total in Words: {data.get('amount_words', '')}")
        elif sheet_name == "Extra Items":
            table = doc.add_table(rows=len(data["items"]) + 1, cols=7)
            table.style = "Table Grid"
            table.autofit = False
            # Approx widths to fill 190mm
            ei_widths_mm = [15, 30, 70, 20, 15, 20, 20]
            for idx, w in enumerate(ei_widths_mm):
                for row in table.rows:
                    row.cells[idx].width = Mm(w)
            headers = ["Serial No.", "Remark", "Description", "Quantity", "Unit", "Rate", "Amount"]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("remark", ""))
                row.cells[2].text = str(item.get("description", ""))
                row.cells[3].text = str(item.get("quantity", ""))
                row.cells[4].text = str(item.get("unit", ""))
                row.cells[5].text = str(item.get("rate", ""))
                row.cells[6].text = str(item.get("amount", ""))
        elif sheet_name == "Deviation Statement":
            # Mirror HTML: 13 columns including Remarks
            table = doc.add_table(rows=len(data["items"]) + 5, cols=13)
            table.style = "Table Grid"
            table.autofit = False
            # Landscape content width = 297 - (2*10) = 277mm
            ds_widths_mm = [6, 90, 9, 9, 9, 9, 9, 9, 9, 9, 9, 9, 50]
            for idx, w in enumerate(ds_widths_mm):
                for row in table.rows:
                    row.cells[idx].width = Mm(w)
            headers = ["ITEM No.", "Description", "Unit", "Qty as per Work Order", "Rate", "Amt as per Work Order Rs.", "Qty Executed", "Amt as per Executed Rs.", "Excess Qty", "Excess Amt Rs.", "Saving Qty", "Saving Amt Rs.", "REMARKS/ REASON."]
            for j, header in enumerate(headers):
                table.rows[0].cells[j].text = header
            for i, item in enumerate(data["items"]):
                row = table.rows[i + 1]
                row.cells[0].text = str(item.get("serial_no", ""))
                row.cells[1].text = str(item.get("description", ""))
                row.cells[2].text = str(item.get("unit", ""))
                row.cells[3].text = str(item.get("qty_wo", ""))
                row.cells[4].text = str(item.get("rate", ""))
                row.cells[5].text = str(item.get("amt_wo", ""))
                row.cells[6].text = str(item.get("qty_bill", ""))
                row.cells[7].text = str(item.get("amt_bill", ""))
                row.cells[8].text = str(item.get("excess_qty", ""))
                row.cells[9].text = str(item.get("excess_amt", ""))
                row.cells[10].text = str(item.get("saving_qty", ""))
                row.cells[11].text = str(item.get("saving_amt", ""))
                row.cells[12].text = str(item.get("remark", ""))
            row = table.rows[-4]
            row.cells[1].text = "Grand Total"
            row.cells[5].text = str(data["summary"].get("work_order_total", ""))
            row.cells[7].text = str(data["summary"].get("executed_total", ""))
            row.cells[9].text = str(data["summary"].get("overall_excess", ""))
            row.cells[11].text = str(data["summary"].get("overall_saving", ""))
            row = table.rows[-3]
            row.cells[1].text = f"Add Tender Premium ({data['summary']['premium'].get('percent', 0):.2%})"
            row.cells[5].text = str(data["summary"].get("tender_premium_f", ""))
            row.cells[7].text = str(data["summary"].get("tender_premium_h", ""))
            row.cells[9].text = str(data["summary"].get("tender_premium_j", ""))
            row.cells[11].text = str(data["summary"].get("tender_premium_l", ""))
            row = table.rows[-2]
            row.cells[1].text = "Grand Total including Tender Premium"
            row.cells[5].text = str(data["summary"].get("grand_total_f", ""))
            row.cells[7].text = str(data["summary"].get("grand_total_h", ""))
            row.cells[9].text = str(data["summary"].get("grand_total_j", ""))
            row.cells[11].text = str(data["summary"].get("grand_total_l", ""))
            row = table.rows[-1]
            net_difference = data["summary"].get("net_difference", 0)
            row.cells[1].text = "Overall Excess" if net_difference > 0 else "Overall Saving"
            row.cells[7].text = str(abs(round(net_difference)))
        elif sheet_name == "Note Sheet":
            for note in data.get("notes", []):
                doc.add_paragraph(str(note))
        doc.save(doc_path)
        _log_debug(f"Finished Word doc for {sheet_name}")
    except Exception as e:
        st.error(f"Error creating Word doc for {sheet_name}: {str(e)}")
        raise

# Streamlit app
uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")
if not uploaded_file:
    _render_landing()
    st.stop()

st.title("Bill Generator")
st.write("Upload an Excel file and enter tender premium details.")
premium_percent = st.number_input("Tender Premium %", min_value=0.0, max_value=100.0, step=0.01)
premium_type = st.selectbox("Premium Type", ["Above", "Below"])

if uploaded_file is not None and st.button("Generate Bill"):
    try:
        xl = pd.ExcelFile(uploaded_file)
        ws_wo = xl.parse("Work Order", header=None)
        ws_bq = xl.parse("Bill Quantity", header=None)
        ws_extra = xl.parse("Extra Items", header=None)

        first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data = process_bill(
            ws_wo, ws_bq, ws_extra, premium_percent, premium_type.lower()
        )

                # Generate note sheet
        try:
            work_order_amount = sum(
                float(ws_wo.iloc[i, 3]) * float(ws_wo.iloc[i, 4])
                for i in range(21, ws_wo.shape[0])
                if pd.notnull(ws_wo.iloc[i, 3]) and pd.notnull(ws_wo.iloc[i, 4])
            )
        except Exception as e:
            st.error(f"Error calculating work_order_amount: {e}")
            work_order_amount = 854678  # Fallback value

        extra_item_amount = first_page_data["totals"].get("extra_items_sum", 0)
        payable_amount = first_page_data["totals"].get("payable", 0)
        note_sheet_data = generate_bill_notes(payable_amount, work_order_amount, extra_item_amount)

        # Define work_order_data from ws_wo or a Work Order sheet
        work_order_data = {
            'agreement_no': ws_wo.iloc[0, 1] if pd.notnull(ws_wo.iloc[0, 1]) else '48/2024-25',
            'name_of_work': ws_wo.iloc[1, 1] if pd.notnull(ws_wo.iloc[1, 1]) else 'Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur',
            'name_of_firm': ws_wo.iloc[2, 1] if pd.notnull(ws_wo.iloc[2, 1]) else 'M/s Seema Electrical Udaipur',
            'date_commencement': ws_wo.iloc[3, 1] if pd.notnull(ws_wo.iloc[3, 1]) else '18/01/2025',
            'date_completion': ws_wo.iloc[4, 1] if pd.notnull(ws_wo.iloc[4, 1]) else '17/04/2025',
            'actual_completion': ws_wo.iloc[5, 1] if pd.notnull(ws_wo.iloc[5, 1]) else '01/03/2025',
            'work_order_amount': str(work_order_amount)
        }

        # Prepare note_sheet_data with VBA-style notes
        percentage_work_done = (float(payable_amount) / float(work_order_amount) * 100) if work_order_amount > 0 else 0
        notes = [
            f"1. The work has been completed {percentage_work_done:.2f}% of the Work Order Amount."
        ]
        if percentage_work_done < 90:
            notes.append("2. The execution of work at final stage is less than 90% of the Work Order Amount, the Requisite Deviation Statement is enclosed to observe check on unuseful expenditure. Approval of the Deviation is having jurisdiction under this office.")
        elif 100 < percentage_work_done <= 105:
            notes.append("2. Requisite Deviation Statement is enclosed. The Overall Excess is less than or equal to 5% and is having approval jurisdiction under this office.")
        elif percentage_work_done > 105:
            notes.append("2. Requisite Deviation Statement is enclosed. The Overall Excess is more than 5% and Approval of the Deviation Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
        delay_days = (datetime.strptime(work_order_data['actual_completion'], '%d/%m/%Y') - datetime.strptime(work_order_data['date_completion'], '%d/%m/%Y')).days
        if delay_days > 0:
            time_allowed = (datetime.strptime(work_order_data['date_completion'], '%d/%m/%Y') - datetime.strptime(work_order_data['date_commencement'], '%d/%m/%Y')).days
            notes.append(f"3. Time allowed for completion of the work was {time_allowed} days. The Work was delayed by {delay_days} days.")
            if delay_days > 0.5 * time_allowed:
                notes.append("4. Approval of the Time Extension Case is required from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
            else:
                notes.append("4. Approval of the Time Extension Case is to be done by this office.")
        else:
            notes.append("3. Work was completed in time.")
        if extra_item_amount > 0:
            extra_item_percentage = (extra_item_amount / float(work_order_amount) * 100) if work_order_amount > 0 else 0
            if extra_item_percentage > 5:
                notes.append(f"4. The amount of Extra items is Rs. {extra_item_amount} which is {extra_item_percentage:.2f}% of the Work Order Amount; exceed 5%, require approval from the Superintending Engineer, PWD Electrical Circle, Udaipur.")
            else:
                notes.append(f"4. The amount of Extra items is Rs. {extra_item_amount} which is {extra_item_percentage:.2f}% of the Work Order Amount; under 5%, approval of the same is to be granted by this office.")
        notes.extend([
            "5. Quality Control (QC) test reports attached.",
            "6. Please peruse above details for necessary decision-making.",
            "",
            "                                Premlata Jain",
            "                               AAO- As Auditor"
        ])

        note_sheet_data = {
            'agreement_no': work_order_data.get('agreement_no', '48/2024-25'),
            'name_of_work': work_order_data.get('name_of_work', 'Electric Repair and MTC work at Govt. Ambedkar hostel Ambamata, Govardhanvilas, Udaipur'),
            'name_of_firm': work_order_data.get('name_of_firm', 'M/s Seema Electrical Udaipur'),
            'date_commencement': work_order_data.get('date_commencement', '18/01/2025'),
            'date_completion': work_order_data.get('date_completion', '17/04/2025'),
            'actual_completion': work_order_data.get('actual_completion', '01/03/2025'),
            'work_order_amount': work_order_data.get('work_order_amount', '854678'),
            'extra_item_amount': extra_item_amount,
            'notes': notes,
            'totals': first_page_data.get('totals', {'payable': str(payable_amount)})
        }
        _log_debug("Note Sheet data prepared")

        # Generate PDFs
        pdf_files = []
        for sheet_name, data, orientation in [
            ("First Page", first_page_data, "portrait"),
            ("Deviation Statement", deviation_data, "landscape"),
            ("Note Sheet", note_sheet_data, "portrait"),
            ("Last Page", last_page_data, "portrait"),
            ("Extra Items", extra_items_data, "portrait"),
        ]:
            pdf_path = os.path.join(get_temp_dir(), f"{sheet_name.replace(' ', '_')}.pdf")
            generate_pdf(sheet_name, data, orientation, pdf_path)
            pdf_files.append(pdf_path)

        # Compile LaTeX templates to a separate output folder
        latex_output_dir = os.path.join(get_temp_dir(), "latex_pdfs")
        latex_pdfs = compile_latex_templates(latex_output_dir)

        # Merge PDFs
        current_date = datetime.now().strftime("%Y%m%d")
        pdf_output = os.path.join(get_temp_dir(), f"BILL_AND_DEVIATION_{current_date}.pdf")
        #############################################################################
        writer = PdfWriter()

        for pdf in pdf_files:
            if os.path.exists(pdf):
                reader = PdfReader(pdf)
                for page in reader.pages:
                    writer.add_page(page)

        if len(writer.pages) > 0:
            with open(pdf_output, "wb") as out_file:
                writer.write(out_file)
        ###########################################################################

        # Generate Word docs
        word_files = []
        for sheet_name, data in [
            ("First Page", first_page_data),
            ("Last Page", last_page_data),
            ("Extra Items", extra_items_data),
            ("Deviation Statement", deviation_data),
            ("Note Sheet", note_sheet_data)
        ]:
            doc_path = os.path.join(get_temp_dir(), f"{sheet_name.replace(' ', '_')}.docx")
            create_word_doc(sheet_name, data, doc_path)
            word_files.append(doc_path)

        # Create ZIP
        zip_path = os.path.join(get_temp_dir(), "bill_output.zip")
        try:
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                if os.path.exists(pdf_output):
                    zipf.write(pdf_output, os.path.basename(pdf_output))
                for word_file in word_files:
                    if os.path.exists(word_file):
                        zipf.write(word_file, os.path.basename(word_file))
                # include individual PDFs and HTMLs for verification
                for sheet_name in ["First Page", "Deviation Statement", "Note Sheet", "Last Page", "Extra Items"]:
                    base = sheet_name.replace(" ", "_")
                    pdf_path = os.path.join(get_temp_dir(), f"{base}.pdf")
                    html_path = os.path.join(get_temp_dir(), f"{base}.html")
                    if os.path.exists(pdf_path):
                        zipf.write(pdf_path, os.path.basename(pdf_path))
                    if os.path.exists(html_path):
                        zipf.write(html_path, os.path.basename(html_path))
                # include LaTeX-compiled PDFs under a separate folder in the ZIP
                for pdf in latex_pdfs:
                    zipf.write(pdf, os.path.join("latex_pdfs", os.path.basename(pdf)))
            with open(zip_path, "rb") as f:
                st.download_button(
                    label="Download Bill Output",
                    data=f,
                    file_name="bill_output.zip",
                    mime="application/zip"
                )
        except Exception as e:
            st.error(f"Error creating ZIP file: {str(e)}")

        # Clean up temporary files
        try:
            temp_dir = get_temp_dir()
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                TEMP_DIR = None
        except Exception as e:
            st.warning(f"Failed to clean up temp directory: {str(e)}")

    except Exception as e:
        st.error(f"Error: {str(e)}")
        _log_traceback()

def run_cli(input_xlsx: str, premium_percent: float, premium_type: str, output_dir: str) -> str:
    import pandas as pd
    os.makedirs(output_dir, exist_ok=True)
    global TEMP_DIR
    TEMP_DIR = output_dir  # For CLI mode, use the specified output directory
    xl = pd.ExcelFile(input_xlsx)
    ws_wo = xl.parse("Work Order", header=None)
    ws_bq = xl.parse("Bill Quantity", header=None)
    ws_extra = xl.parse("Extra Items", header=None)
    first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data = process_bill(
        ws_wo, ws_bq, ws_extra, premium_percent, premium_type.lower()
    )
    # Compute note_sheet_data same as UI flow
    try:
        work_order_amount = sum(
            float(ws_wo.iloc[i, 3]) * float(ws_wo.iloc[i, 4])
            for i in range(21, ws_wo.shape[0])
            if pd.notnull(ws_wo.iloc[i, 3]) and pd.notnull(ws_wo.iloc[i, 4])
        )
    except Exception:
        work_order_amount = 0
    extra_item_amount = first_page_data["totals"].get("extra_items_sum", 0)
    payable_amount = first_page_data["totals"].get("payable", 0)
    ns = generate_bill_notes(payable_amount, work_order_amount, extra_item_amount)
    # Work order metadata (fallbacks similar to UI)
    work_order_data = {
        'agreement_no': ws_wo.iloc[0, 1] if pd.notnull(ws_wo.iloc[0, 1]) else '48/2024-25',
        'name_of_work': ws_wo.iloc[1, 1] if pd.notnull(ws_wo.iloc[1, 1]) else 'N/A',
        'name_of_firm': ws_wo.iloc[2, 1] if pd.notnull(ws_wo.iloc[2, 1]) else 'N/A',
        'date_commencement': ws_wo.iloc[3, 1] if pd.notnull(ws_wo.iloc[3, 1]) else 'N/A',
        'date_completion': ws_wo.iloc[4, 1] if pd.notnull(ws_wo.iloc[4, 1]) else 'N/A',
        'actual_completion': ws_wo.iloc[5, 1] if pd.notnull(ws_wo.iloc[5, 1]) else 'N/A',
        'work_order_amount': str(work_order_amount)
    }
    note_sheet_data = {
        **note_sheet_data,
        **ns,
        'agreement_no': work_order_data['agreement_no'],
        'name_of_work': work_order_data['name_of_work'],
        'name_of_firm': work_order_data['name_of_firm'],
        'date_commencement': work_order_data['date_commencement'],
        'date_completion': work_order_data['date_completion'],
        'actual_completion': work_order_data['actual_completion'],
        'work_order_amount': work_order_data['work_order_amount'],
        'extra_item_amount': extra_item_amount,
        'totals': first_page_data.get('totals', {'payable': str(payable_amount)})
    }

    # Generate PDFs and DOCX like UI flow
    pdf_files = []
    for sheet_name, data, orientation in [
        ("First Page", first_page_data, "portrait"),
        ("Deviation Statement", deviation_data, "landscape"),
        ("Note Sheet", note_sheet_data, "portrait"),
        ("Last Page", last_page_data, "portrait"),
        ("Extra Items", extra_items_data, "portrait"),
    ]:
        pdf_path = os.path.join(output_dir, f"{sheet_name.replace(' ', '_')}.pdf")
        generate_pdf(sheet_name, data, orientation, pdf_path)
        if os.path.exists(pdf_path):
            pdf_files.append(pdf_path)

    writer = PdfWriter()
    for pdf in pdf_files:
        if os.path.exists(pdf):
            reader = PdfReader(pdf)
            for page in reader.pages:
                writer.add_page(page)
    merged_pdf = os.path.join(output_dir, "BILL_AND_DEVIATION.pdf")
    if len(writer.pages) > 0:
        with open(merged_pdf, "wb") as f:
            writer.write(f)

    word_files = []
    for sheet_name, data in [
        ("First Page", first_page_data),
        ("Last Page", last_page_data),
        ("Extra Items", extra_items_data),
        ("Deviation Statement", deviation_data),
        ("Note Sheet", note_sheet_data)
    ]:
        doc_path = os.path.join(output_dir, f"{sheet_name.replace(' ', '_')}.docx")
        create_word_doc(sheet_name, data, doc_path)
        word_files.append(doc_path)

    # Compile LaTeX
    latex_output_dir = os.path.join(output_dir, "latex_pdfs")
    compile_latex_templates(latex_output_dir)

    # Create zip
    zip_path = os.path.join(output_dir, "bill_output.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        if os.path.exists(merged_pdf):
            zipf.write(merged_pdf, os.path.basename(merged_pdf))
        for wf in word_files:
            if os.path.exists(wf):
                zipf.write(wf, os.path.basename(wf))
        for pf in pdf_files:
            if os.path.exists(pf):
                zipf.write(pf, os.path.basename(pf))
        if os.path.isdir(latex_output_dir):
            for name in os.listdir(latex_output_dir):
                if name.lower().endswith('.pdf'):
                    zipf.write(os.path.join(latex_output_dir, name), os.path.join('latex_pdfs', name))
    return zip_path

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Bill PDFs programmatically")
    parser.add_argument("input", help="Path to input Excel (.xlsx)")
    parser.add_argument("--premium-percent", type=float, default=0.0)
    parser.add_argument("--premium-type", choices=["add", "deduct"], default="add")
    parser.add_argument("--out", default=os.path.join(os.getcwd(), "output"))
    args = parser.parse_args()
    out_zip = run_cli(args.input, args.premium_percent, args.premium_type, args.out)
    print(out_zip)
